from sklearn.model_selection import train_test_split
import docx

import pandas as pd
from sklearn import tree
from sklearn.tree import DecisionTreeClassifier
from sklearn.metrics import accuracy_score
import matplotlib.pyplot as plot
from statistics import mean

dataset = pd.read_csv(
    'BankNote_Authentication.csv')


x = dataset.iloc[:, [0, 1, 2, 3]]
y = dataset.iloc[:, [4]]


def DT(RS, TS):
    # splitting data
    x_train, x_test, y_train, y_test = train_test_split(
        x, y, random_state=RS, train_size=TS)

    # create DecisionTreeClassifier obj and fitting data
    clf = DecisionTreeClassifier()
    clf.fit(x_train, y_train)

    # make a prediction and calculate accuracy (train - test)
    y_pred_test = clf.predict(x_test)
    y_pred_train = clf.predict(x_train)
    #print("Train data accuracy:",accuracy_score(y_train,y_pred_train))
    #print("Test data accuracy:",accuracy_score( y_test, y_pred_test))

    # list of all accuracy values
    acc = accuracy_score(y_test, y_pred_test)
    # number of nodes node  (tree size)
    nodes_num = clf.tree_.node_count
    # list of all number of nodes
    return acc, nodes_num


rs = [1, 50, 15, 100, 20]
ts = [0.3, 0.4, 0.5, 0.6, 0.7]
All_accuracy = []
All_nodes = []
minAcc = []
maxAcc = []
meanAcc = []

minSize = []
maxSize = []
meanSize = []

for i in ts:
    for j in rs:
        accuracy, nodes = DT(j, i)
        All_accuracy.append(accuracy)
        All_nodes.append(nodes)

     # calculate mean, maximum and minimum accuracy at each training set_size.
    minAcc.append(min(All_accuracy))
    maxAcc.append(max(All_accuracy))
    meanAcc.append(sum(All_accuracy)/len(All_accuracy))

    # measure the mean, max and min tree size.
    minSize.append(min(All_nodes))
    maxSize.append(max(All_nodes))
    meanSize.append(sum(All_nodes)/len(All_nodes))

    All_accuracy = []
    All_nodes = []


def listToString(s):

    # initialize an empty string
    str1 = ""

    # traverse in the string
    for ele in s:
        str1 += str(ele)
        str1 += ' '

    # return string
    return str1


doc = docx.Document()
# create a document builder object

doc.add_paragraph('mean accuracy :')
doc.add_paragraph('['+listToString(meanAcc)+']'+"\n"+"\n")
doc.add_paragraph('minimum accuracy : ')
doc.add_paragraph('['+listToString(minAcc)+']'+"\n"+"\n")
doc.add_paragraph('maximim of the tree size : ')
doc.add_paragraph('['+listToString(maxAcc)+']'+"\n"+"\n")
doc.add_paragraph('mean of the tree size : ')
doc.add_paragraph('['+listToString(meanSize)+']'+"\n"+"\n")
doc.add_paragraph('minimum tree size : ')
doc.add_paragraph('['+listToString(minSize)+']' + "\n"+"\n")
doc.add_paragraph('maximum  tree size :')
doc.add_paragraph('['+listToString(maxSize)+']' + "\n"+"\n")
trainSetSize = [.3*1372, 1372*.4, 1372 * .5, 1372 * .6, 1372 * .7]
plot.plot(maxAcc, trainSetSize)
plot.savefig("output.jpg")
doc.add_picture('output.jpg')

plot.plot(minAcc, trainSetSize)

plot.savefig("output.jpg")

doc.add_picture('output.jpg')
plot.plot(meanAcc, trainSetSize)

plot.savefig("output.jpg")
doc.add_picture('output.jpg')
plot.plot(minSize, trainSetSize)
plot.savefig("output.jpg")
doc.add_picture('output.jpg')
plot.plot(maxSize, trainSetSize)
plot.savefig("output.jpg")
doc.add_picture('output.jpg')
plot.plot(meanSize, trainSetSize)


doc.save("out.docx")